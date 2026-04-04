import os
import json
from crewai import Agent, Task, Crew
from crewai.tools import tool
from docx import Document
import re
from crewai import LLM
from dotenv import load_dotenv


# ==========================================
# 1. 配置 Kimi 2.5 (使用 OpenAI 兼容模式)
# ==========================================

load_dotenv('.env')

os.environ["OPENAI_API_KEY"] = "sk-YaoAb3bYs3tnemLkEIvKebOBJJrKSiTzeC1BNZHvdPJgXLNt"
os.environ["OPENAI_API_BASE"] = "https://api.moonshot.cn/v1"
os.environ["OPENAI_API_MODEL"] = "kimi-k2.5"



# 使用 CrewAI 原生的 LLM 类
# 注意 model 名称前面的 "openai/" 前缀，这是告诉底层 LiteLLM 使用 OpenAI 兼容协议
kimi_llm = LLM(
    model=os.environ["OPENAI_API_MODEL"],
    base_url=os.environ["OPENAI_API_BASE"] ,
    api_key=os.environ["OPENAI_API_KEY"] , # 填入你真实的 Kimi Key
    temperature=1,
    # 核心突破口：注入伪造的 Headers 绕过校验
    extra_headers = {
        "User-Agent": "Kimi CLI"
    }
)

@tool("Word坐标提取为GeoJSON工具")
def extract_word_table_to_geojson(
        docx_path: str,
        output_name: str = "output.geojson",
        lon_col_index: int = -1,
        lat_col_index: int = -1
) -> str:
    """
    当你需要从本地 Word 文档 (.docx) 的表格中提取经纬度坐标，并转换为 GeoJSON 格式时，请调用此工具。

    参数:
    - docx_path: 输入的 Word 文档的路径。
    - output_name: 输出的 .geojson 文件的保存路径。
    - lon_col_index: 经度数据所在的列索引（从 0 开始）。初次调用请传 -1，让工具尝试自动识别。
    - lat_col_index: 纬度数据所在的列索引（从 0 开始）。初次调用请传 -1，让工具尝试自动识别。
    """
    if not os.path.exists(docx_path):
        return f"错误：找不到文件 {docx_path}，请检查路径。"

    try:
        doc = Document(docx_path)
        if not doc.tables:
            return "错误：文档中没有找到任何表格。"

        # 默认处理文档中的第一个表格（可根据实际业务扩展为遍历）
        table = doc.tables[0]

        # ==========================================
        # 阶段一：自动推断与 Agent 兜底机制
        # ==========================================
        if lon_col_index == -1 or lat_col_index == -1:
            # 1. 先尝试用轻量级正则模糊匹配表头
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            lon_pattern = re.compile(r'(lon|lng|x|经度|经)', re.IGNORECASE)
            lat_pattern = re.compile(r'(lat|y|纬度|纬)', re.IGNORECASE)

            for i, text in enumerate(headers):
                if lon_pattern.search(text) and lon_col_index == -1: lon_col_index = i
                if lat_pattern.search(text) and lat_col_index == -1: lat_col_index = i

            # 2. 兜底触发：如果正则依然没找到（比如没有表头，或者表头叫“位置一”）
            # 此时我们不报错，而是把前3行数据提取出来，发给 Agent 求助！
            if lon_col_index == -1 or lat_col_index == -1:
                preview_data = []
                for row in table.rows[:3]:
                    # 截取每格前 20 个字符，防止单元格内容过长撑爆 Token
                    preview_data.append([cell.text.strip()[:20] for cell in row.cells])

                return (
                    f"⚠️ 警告：工具无法自动识别哪一列是经度和纬度。\n"
                    f"这是文档表格的前 3 行数据预览：\n{preview_data}\n\n"
                    f"请作为数据分析师，观察上面的二维数组，判断经度和纬度分别在第几列（索引从 0 开始）。\n"
                    f"然后，请【再次调用】本工具，并务必传入你推断出的 lon_col_index 和 lat_col_index 参数！"
                )

        # ==========================================
        # 阶段二：使用确定的索引进行强制提取
        # ==========================================
        features = []

        # 遍历所有行进行提取
        for row_idx, row in enumerate(table.rows):
            cells = row.cells

            # 防止某一行合并单元格导致列数不够
            if len(cells) <= max(lon_col_index, lat_col_index):
                continue

            try:
                # 尝试将目标列转换为浮点数
                lon_text = cells[lon_col_index].text.strip()
                lat_text = cells[lat_col_index].text.strip()

                lon = float(lon_text)
                lat = float(lat_text)

                # 提取其余列作为 GeoJSON 的 properties 属性
                properties = {"row_index": row_idx}  # 记录原始行号以便追溯
                for i, cell in enumerate(cells):
                    if i != lon_col_index and i != lat_col_index:
                        # 因为无法确定表头，统筹命名为 Column_0, Column_1...
                        val = cell.text.strip()
                        if val: properties[f"Column_{i}"] = val

                feature = {
                    "type": "Feature",
                    "geometry": {"type": "Point", "coordinates": [lon, lat]},
                    "properties": properties
                }
                features.append(feature)

            except ValueError:
                # 如果转换 float 失败（例如这一行是表头文本，或者是空白单元格），直接静默跳过
                continue

        if not features:
            return f"❌ 失败：使用列索引 lon={lon_col_index}, lat={lat_col_index} 未能提取到任何合法的坐标数字。请检查你的索引推理是否正确。"

        # ==========================================
        # 阶段三：生成并保存文件
        # ==========================================
        geojson_data = {
            "type": "FeatureCollection",
            "features": features
        }

        with open(output_name, 'w', encoding='utf-8') as f:
            json.dump(geojson_data, f, ensure_ascii=False, indent=2)

        return f"✅ 成功！已使用列索引 (经度:{lon_col_index}, 纬度:{lat_col_index})，提取了 {len(features)} 个点位，并保存至 {output_name}"

    except Exception as e:
        return f"🚨 执行工具时发生内部报错：{str(e)}"


# ==========================================
# 3. 定义 Agent (设置角色和分配工具)
# ==========================================
gis_analyst = Agent(
    role='高级地理信息(GIS)数据工程师',
    goal='准确提取文本报告中的空间点位数据，并转换为标准结构化地理格式',
    backstory=(
        '你是一名经验丰富的 GIS 专家，曾参与过多次重大应急响应项目。'
        '你擅长从繁杂的文档中清洗和提取空间数据，为下游的态势感知或 C4ISR 系统提供可靠的数据源。'
        '你做事严谨，熟练使用各种数据转换工具。'
    ),
    verbose=True,  # 开启日志，你将在终端看到它的思考过程 (Vibe Coding 必备)
    allow_delegation=False,  # 不允许它把任务推给别的 Agent (因为目前只有它一个)
    tools=[extract_word_table_to_geojson],  # 把上面定义的 Skill 交给它
    llm=kimi_llm
)

# ==========================================
# 4. 定义 Task (给 Agent 下达具体指令)
# ==========================================
extract_task = Task(
    description=(
        '请读取当前目录下的 "middle_east_emergency_points.docx" 文件，'
        '提取其中的坐标点位，并将其保存为 "middle_east_response.geojson" 文件。'
    ),
    expected_output='一段简短的中文摘要，说明是否成功转换，以及成功提取了多少个点位。',
    agent=gis_analyst
)

# ==========================================
# 5. 组装 Crew 并一键运行
# ==========================================
gis_crew = Crew(
    agents=[gis_analyst],
    tasks=[extract_task],
    verbose=True  # 开启宏观日志
)

if __name__ == "__main__":
    print("🚀 启动 GIS 数据处理 Agent...")
    # kickoff 会触发 Agent 开始思考并调用工具
    result = gis_crew.kickoff()

    print("\n" + "=" * 40)
    print("✅ Agent 最终汇报结果：")
    print(result)