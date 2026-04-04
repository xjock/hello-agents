from docx import Document

def create_test_word_doc(file_name="middle_east_emergency_points.docx"):
    # 创建一个新的 Word 文档
    doc = Document()
    doc.add_heading('中东地区应急响应点位坐标（测试数据）', 0)

    # 添加一段说明文字
    doc.add_paragraph('以下表格包含了本次应急响应的关键坐标点位，请提取并转换为 GeoJSON 格式。')

    # 创建一个 5 行 4 列的表格（1行表头 + 4行数据）
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid' # 添加表格边框，使其在 Word 中看起来更清晰

    # 设置表头 (故意使用中英文混杂，测试正则匹配)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '点位名称'
    hdr_cells[1].text = '经度 (Lon)'
    hdr_cells[2].text = '纬度 (Lat)'
    hdr_cells[3].text = '现场情况备注'

    # 准备测试数据 (模拟一些中东区域的真实经纬度范围)
    test_data = [
        ('大马士革集结区', '36.2765', '33.5138', '主要物资中转站，目前安全'),
        ('阿勒颇临时医疗点', '37.1342', '36.2021', '急需外科手术设备'),
        ('霍姆斯检查站', '36.7233', '34.7355', '交通拥堵，有人员滞留')
    ]

    # 将正常数据填入表格
    for item in test_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]

    # 故意制造一行“脏数据”（坐标缺失），测试提取函数的 try-except 容错能力
    dirty_row = table.add_row().cells
    dirty_row[0].text = '未知通讯盲区'
    dirty_row[1].text = '待查'    # 无法转为 float
    dirty_row[2].text = '待查'    # 无法转为 float
    dirty_row[3].text = '坐标丢失，等待前方人员携带设备确认'

    # 保存文档
    doc.save(file_name)
    print(f"✅ 测试文档生成成功：{file_name}，快去让 Agent 提取吧！")

if __name__ == "__main__":
    create_test_word_doc()